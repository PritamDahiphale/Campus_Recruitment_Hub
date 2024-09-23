
from flask import Flask, request, redirect, url_for, session, flash, send_file
from flask.templating import render_template
import sqlite3  
import re

# Initialize app
app = Flask(__name__)

app.secret_key = 'your secret key'

app.debug=True


'''<----------------------------All routes for home page--------------------------------'''
@app.route("/")
def home():
	return render_template("index.html")

@app.route("/about")
def about():
	return render_template("about.html")

@app.route("/login")
def login():
	return render_template("login.html")

@app.route("/forgot", methods =['GET', 'POST'])
def forgot():
	if request.method == 'POST' and 'username' in request.form:
		username = request.form['username']

		con = sqlite3.connect("users.db")  
		cur=con.cursor()
		cur.execute('SELECT * FROM Users WHERE username =?',[username])
		account = cur.fetchone()
		if account:
			flash('Correct Username! Now you can reset your password.')
			return render_template('reset_pass.html',username=username)
		else:
			flash("Try again! Enter correct username.")
			return render_template('login.html')

@app.route("/reset_pass_db", methods =['GET', 'POST'])
def reset_pass_db():
	if request.method == 'POST' and 'password' in request.form and 'confirm_password' in request.form:
		username = request.form['username']
		password = request.form['password']
		confirm_password = request.form['confirm_password']

		if(password==confirm_password):
			con = sqlite3.connect("Users.db")  
			cur=con.cursor()
			cur.execute('UPDATE Users SET password=? WHERE username=?', [password, username])
			con.commit()
			flash('Password resetted successfully!! Try to login again.')
			return render_template("login.html")
		else:
			flash("Both password fields doesn't match!!!")
			return render_template('reset_pass.html')

@app.route("/signup")
def signup():
	return render_template("signup.html")

@app.route("/contact")
def contact():
	return render_template("contact.html")

@app.route("/notice_main")
def notice_main():
	title = "Notices"
	con = sqlite3.connect("notice.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from notice")   
	notices = cur.fetchall()  
	totalnotices=len(notices) 
	return render_template("notice_main.html", title=title,notices=notices,totalnotices=totalnotices)
	

# Login and signup 
@app.route('/loginlog', methods =['GET', 'POST'])
def loginlog():
	if request.method == 'POST' and 'username' in request.form and 'password' in request.form:
			username = request.form['username']
			password = request.form['password']
			
			con = sqlite3.connect("users.db")  
			cur=con.cursor()
			cur.execute('SELECT * FROM Users WHERE username = ? AND password = ?', (username, password))
			account = cur.fetchall()

			con2 = sqlite3.connect("admin.db")
			cur2=con2.cursor()
			cur2.execute('SELECT * FROM Admin WHERE username = ? AND password = ?', (username,password))
			adminn = cur2.fetchone()

			if account:
				session['loggedin'] = True
				session['username'] = username
				flash('Logged in successfully !')
				return render_template('student_dash.html')

			elif adminn:
				session['loggedin'] = True
				session['username'] = username
				flash('Logged in successfully !')
				return redirect(url_for('dashboard'))

			else:
				flash('Incorrect username / password! Sorry!! We cannot let you login!!')
	return render_template('login.html')
		


@app.route('/logout')
def logout():
	session.pop('loggedin', None)
	session.pop('id', None)
	session.pop('username', None)
	flash("Successfully logout!!")
	return redirect(url_for('login'))

@app.route('/signupsign', methods =['GET', 'POST'])
def signupsign():
	
	if request.method == 'POST' and 'username' in request.form and 'password' in request.form and 'email' in request.form :
		username = request.form['username']
		password = request.form['password']
		email = request.form['email']

		con = sqlite3.connect("Users.db")  
		cur=con.cursor()
		cur.execute('SELECT * FROM Users WHERE username =?',[username])
		account = cur.fetchone()
		
		if account:
			flash('Account already exists !')
		elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
			flash('Invalid email address !')
		elif not re.match(r'[A-Za-z0-9]+', username):
			flash('Username must contain only characters and numbers !')
		elif not re.match(r'[A-Za-z0-9]+', password):
			flash('Password must contain characters and numbers!')
		elif not username or not password or not email:
			flash('Please fill out the form !')
		else:
			cur.execute('INSERT INTO Users (username, password, email) VALUES (?,?,?)', (username, password, email))
			con.commit()
			flash('You have successfully registered !')
	elif request.method == 'POST':
		flash('Please fill out the form !')
	return render_template('signup.html')

'''<----------------------------All routes for admin page--------------------------------'''

@app.route("/dashboard")
def dashboard():
	title = "Admin Dashboard"
	return render_template("dashboard.html", title=title)

@app.route("/viewStudents")
def viewStudents():
	con = sqlite3.connect("student.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from student")   
	details = cur.fetchall()  
	totalstud=len(details)
	return render_template("viewStudents.html", details=details,totalstud=totalstud)


@app.route("/company")
def company():
	con = sqlite3.connect("Company.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from Company")   
	companies = cur.fetchall()  
	totalcomp=len(companies)
	return render_template("company.html",companies=companies, totalcomp=totalcomp)


@app.route("/insertcompany",methods = ["POST","GET"])
def insertcompany():
	
	if request.method == "POST":  
		try:  
			comp_name = request.form["comp_name"] 
			comp_joined_date = request.form["comp_joined_date"] 
			comp_turnover = request.form["comp_turnover"] 
			comp_website = request.form["comp_website"]  
			 
			with sqlite3.connect("Company.db") as con: 
				 
				cur = con.cursor()   
				# exist = cur.fetchone()
				# if(exist is None):
				# 	cur.execute("drop table Company")
				# 	con.execute('''create table Company (comp_id INTEGER PRIMARY KEY AUTOINCREMENT, 
				# 				comp_name TEXT UNIQUE NOT NULL,comp_joined_date DATE NOT NULL,comp_turnover REAL,
				# 				comp_website TEXT NOT NULL)''') 
				# 	cur.execute("INSERT into Company (comp_name,comp_joined_date,comp_turnover,comp_website) values (?,?,?,?)",(comp_name,comp_joined_date,comp_turnover,comp_website))  
				# 	con.commit()  
				# 	flash("Company successfully Added")  
				# else:
				cur.execute("INSERT into Company (comp_name,comp_joined_date,comp_turnover,comp_website) values (?,?,?,?)",(comp_name,comp_joined_date,comp_turnover,comp_website))  
				con.commit()  
				flash("Company added successfully!")
		except:  
			con.rollback()
			flash("We can not add the Company to the list")
		finally:  
			con = sqlite3.connect("Company.db")  
			con.row_factory = sqlite3.Row  
			cur = con.cursor()  
			cur.execute("select * from Company")   
			companies = cur.fetchall()  
			totalcomp=len(companies)
			return render_template("company.html",companies=companies,totalcomp=totalcomp)
			

@app.route("/updatecomp",methods=('GET', 'POST'))
def updatecomp():
	
	if request.method=='POST':
		try:
			comp_id=request.form["comp_id"]
			comp_name=request.form["comp_name"] 
			comp_joined_date=request.form['comp_joined_date']
			comp_turnover=request.form['comp_turnover']
			comp_website=request.form['comp_website']
			
			with sqlite3.connect("Company.db") as con: 
					
				cur = con.cursor()   
				cur.execute("UPDATE Company SET comp_name=?, comp_joined_date=?, comp_turnover=?, comp_website=? WHERE comp_id=?",(comp_name,comp_joined_date,comp_turnover,comp_website,comp_id))  
				con.commit() 
				flash("Company updated successfully!") 
			
		except:  
			flash("We can not update the Company.")
			return render_template("company.html", companies=companies)

		finally:  
			con = sqlite3.connect("Company.db")  
			con.row_factory = sqlite3.Row  
			cur = con.cursor()  
			cur.execute("select * from Company")   
			companies = cur.fetchall() 
			totalcomp=len(companies) 
			return render_template("company.html",companies=companies,totalcomp=totalcomp)
			
		
@app.route("/deletecomp/<int:del_id>",methods = ["POST"])  
def deletecomp(del_id):  
	
	with sqlite3.connect("Company.db") as con:  
		try:  
			cur = con.cursor()   
			cur.execute("delete from Company where comp_id = ?",(del_id,)) 
			con.commit()
			flash("Company deleted successfully.")
			return redirect('/list')   
		except:  
			flash("Company cannot be deleted.")
			return redirect('/list')
    
@app.route('/list')
def list():
	con = sqlite3.connect("Company.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from Company")   
	companies = cur.fetchall()  
	totalcomp=len(companies)
	return render_template("company.html",companies=companies,totalcomp=totalcomp)	

@app.route("/events")
def events():
	return redirect('/list2')

@app.route("/insertevent",methods = ["POST","GET"])
def insertevent():
	
	if request.method == "POST":  
		try:  
			event_title=request.form['event_title']
			event_name = request.form["event_name"] 
			event_date = request.form["event_date"] 
			 
			with sqlite3.connect("events.db") as con:  
				cur = con.cursor()   
				cur.execute("INSERT into event (event_title,event_name,event_date) values (?,?,?)",(event_title,event_name,event_date))  
				con.commit()  
				flash("Event added successfully!!") 
				return redirect('/list2')
		except:  
			con.rollback()
			flash("We can not add the Event to the list. Try again!!") 
		finally:  
			return redirect('/list2')


@app.route("/update_event",methods=('GET', 'POST'))
def update_event():
	
	if request.method=='POST':
		try:
			event_id=request.form['event_id']
			event_title=request.form['event_title']
			event_name = request.form["event_name"] 
			event_date = request.form["event_date"] 
			
			with sqlite3.connect("events.db") as con: 
					
				cur = con.cursor()   
				cur.execute("UPDATE event SET event_title=?, event_name=?, event_date=? WHERE event_id=?",(event_title,event_name,event_date,event_id))  
				con.commit() 
				flash("Event updated successfully!") 
			
		except:  
			flash("We can not update the Event.")
			return redirect('/list2')

		finally:  
			return redirect('/list2')
			
		
@app.route("/delete_event/<int:del_id>",methods = ["POST"])  
def delete_event(del_id):  
	
	with sqlite3.connect("events.db") as con:  
		try:  
			cur = con.cursor()   
			cur.execute("delete from event where event_id = ?",(del_id,)) 
			con.commit()
			flash("Event deleted successfully.")
			return redirect('/list2')   
		except:  
			flash("Event cannot be deleted.")
			return redirect('/list2')
    
@app.route('/list2')
def list2():
	con = sqlite3.connect("events.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from event")   
	events = cur.fetchall()  
	totalevents=len(events) 
	return render_template("event.html",events=events,totalevents=totalevents)	


@app.route("/notices")
def notices():
	return redirect('/list3')  

@app.route("/insertnotice",methods = ["POST","GET"])
def insertnotice():
	
	if request.method == "POST":  
		try:  
			notice_title=request.form['notice_title']
			notice_name = request.form["notice_name"] 
			notice_date = request.form["notice_date"] 
			 
			with sqlite3.connect("notice.db") as con:  
				cur = con.cursor()   
				cur.execute("INSERT into notice (notice_title,notice_name,notice_date) values (?,?,?)",(notice_title,notice_name,notice_date))  
				con.commit()  
				flash("Notice added successfully!!") 
				return redirect('/list3')  
		except:  
			con.rollback()
			flash("We can not add the notice to the list. Try again!!") 
		finally:  
			return redirect('/list3')  


@app.route("/update_notice",methods=('GET', 'POST'))
def update_notice():
	
	if request.method=='POST':
		try:
			notice_id=request.form['notice_id']
			notice_title=request.form['notice_title']
			notice_name = request.form["notice_name"] 
			notice_date = request.form["notice_date"]  
			
			with sqlite3.connect("notice.db") as con: 
					
				cur = con.cursor()   
				cur.execute("UPDATE notice SET notice_title=?, notice_name=?, notice_date=? WHERE notice_id=?",(notice_title,notice_name,notice_date,notice_id))  
				con.commit() 
				flash("Notice updated successfully!") 
			
		except:  
			flash("We can not update the Notice.")
			return redirect('/list3')

		finally:  
			return redirect('/list3')
			
		
@app.route("/delete_notice/<int:del_id>",methods = ["POST"])  
def delete_notice(del_id):  
	
	with sqlite3.connect("notice.db") as con:  
		try:  
			cur = con.cursor()   
			cur.execute("delete from notice where notice_id = ?",(del_id,)) 
			con.commit()
			flash("Notice deleted successfully.")
			return redirect('/list3')   
		except:  
			flash("Notice cannot be deleted.")
			return redirect('/list3')
    
@app.route('/list3')
def list3():
	con = sqlite3.connect("notice.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from notice")   
	notices = cur.fetchall()  
	totalnotices=len(notices) 
	return render_template("notice.html",notices=notices,totalnotices=totalnotices)	


@app.route("/stat_up")
def stat_up():
	return render_template("statistic.html")


@app.route("/fileupload")
def fileupload():
	return render_template("upload_file.html")

@app.route("/recruit")
def recruit():
	return redirect('/list4') 


@app.route("/insertdata",methods = ["POST","GET"])
def insertdata():
	
	if request.method == "POST":  
		try:  
			comp_name = request.form["comp_name"] 
			recruiter_name = request.form["recruiter_name"] 
			recruit_date = request.form["recruit_date"] 
			selected_stud=request.form["selected_stud"]
			comp_website = request.form["comp_website"]  
			 
			with sqlite3.connect("recruit.db") as con: 
				 
				cur = con.cursor()   
				# exist = cur.fetchone()
				# if(exist is None):
				# 	cur.execute("drop table recruit")
				# 	con.execute('''create table recruit (comp_id INTEGER PRIMARY KEY AUTOINCREMENT, 
				# 				comp_name TEXT UNIQUE NOT NULL,recruiter_name TEXT NOT NULL, recruit_date DATE NOT NULL, selected_stud INTEGER NOT NULL,
				# 				comp_website TEXT NOT NULL)''') 
				# 	cur.execute("INSERT into recruit (comp_name,recruiter_name,recruit_date,selected_stud,comp_website) values (?,?,?,?,?)",(comp_name,recruiter_name,recruit_date,selected_stud,comp_website))  
				# 	con.commit()  
				# 	flash("Data added successfully!")  
				# else:
				cur.execute("INSERT into recruit (comp_name,recruiter_name,recruit_date,selected_stud,comp_website) values (?,?,?,?,?)",(comp_name,recruiter_name,recruit_date,selected_stud,comp_website))  
				con.commit()  
				flash("Data added successfully!")
		except:  
			con.rollback()
			flash("We can not add the data to the list")
		finally:  
			return redirect('/list4') 


 
@app.route("/updatedata",methods=('GET', 'POST'))
def updatedata():
	
	if request.method=='POST':
		try:
			comp_id=request.form['comp_id']
			comp_name = request.form["comp_name"] 
			recruiter_name = request.form["recruiter_name"] 
			recruit_date = request.form["recruit_date"] 
			selected_stud=request.form["selected_stud"]
			comp_website = request.form["comp_website"] 
			
			with sqlite3.connect("recruit.db") as con: 
					
				cur = con.cursor()   
				cur.execute("UPDATE recruit SET comp_name=?, recruiter_name=?, recruit_date=?, selected_stud=?, comp_website=? WHERE comp_id=?",(comp_name,recruiter_name,recruit_date,selected_stud,comp_website,comp_id))  
				con.commit() 
				flash("Data updated successfully!") 
			
		except:  
			flash("We can not update the data.")
			return redirect('/list4')

		finally:  
			return redirect('/list4') 
			
		
@app.route("/deletedata/<int:del_id>",methods = ["POST"])  
def deletedata(del_id):  
	
	with sqlite3.connect("recruit.db") as con:  
		try:  
			cur = con.cursor()   
			cur.execute("delete from recruit where comp_id = ?",(del_id,)) 
			con.commit()
			flash("Data deleted successfully.")
			return redirect('/list4')   
		except:  
			flash("Data cannot be deleted.")
			return redirect('/list4')
    
@app.route('/list4')
def list4():
	con = sqlite3.connect("recruit.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from recruit")   
	recruiters = cur.fetchall()
	totalrecruiters=len(recruiters)  
	return render_template("recruit.html",recruiters=recruiters,totalrecruiters=totalrecruiters)

@app.route("/changepass")
def changepass():
	return render_template("changepass.html")


@app.route("/change_pass_db",methods = ["POST"])
def change_pass_db():
	if request.method=='POST':
		try:
			new_pass=request.form['new_pass']
			username = request.form["username"]  
			
			with sqlite3.connect("admin.db") as con: 
					
				cur = con.cursor()   
				cur.execute('UPDATE Admin SET password=? WHERE username=?', [new_pass, username])
				con.commit()
				flash('Password resetted successfully!!')
			
		except:  
			flash("We can not reset the password.")
			return render_template("changepass.html")

		finally:  
			return render_template("changepass.html")



'''<----------------------------All routes for student page--------------------------------'''


@app.route("/stud_dash")
def stud_dash():
	title = "Student Dashboard"
	return render_template("student_dash.html", title=title)
	
@app.route("/add_detail")
def add_detail():
	return redirect('/list5')

@app.route("/view_detail",methods = ["POST","GET"])
def view_detail():
	if request.method == "POST":  
		
		stud_enroll=request.form['stud_enroll']

		con = sqlite3.connect("student.db")  
		con.row_factory = sqlite3.Row  
		cur = con.cursor()  
		cur.execute("select * from student")   
		details = cur.fetchone()   
		cur.execute("select * from student where stud_enroll=?",[stud_enroll])   
		details_of_each = cur.fetchall()   
		return render_template("add_detail.html",details=details,details_of_each=details_of_each)

@app.route("/insertdetail",methods = ["POST","GET"])
def insertdetail():
	
	if request.method == "POST":  
		try:
			stud_enroll=request.form['stud_enroll']
			stud_name=request.form['stud_name']
			birth_date = request.form["birth_date"] 
			stud_gender = request.form["stud_gender"] 
			stud_address = request.form["stud_address"] 
			stud_branch = request.form["stud_branch"] 
			stud_branch_code = request.form["stud_branch_code"] 
			stud_marks=request.form["stud_marks"]
			stud_back_sub = request.form["stud_back_sub"] 
			stud_in_comp=request.form["stud_in_comp"] 
			
			 
			with sqlite3.connect("student.db") as con:  
				cur = con.cursor()   
				cur.execute("INSERT into student (stud_enroll,stud_name,birth_date,stud_gender,stud_address,stud_branch,stud_branch_code,stud_marks,stud_back_sub,stud_in_comp) values (?,?,?,?,?,?,?,?,?,?)",(stud_enroll,stud_name,birth_date,stud_gender,stud_address,stud_branch,stud_branch_code,stud_marks,stud_back_sub,stud_in_comp))  
				con.commit()  
				flash("Information added successfully!!") 
				return redirect('/list5')
		except:  
			con.rollback()
			flash("We can not add the info.Try again!!")
			return redirect('/list5')
		finally:  
			return redirect('/list5')

@app.route("/update_details",methods=('GET', 'POST'))
def update_details():
	if request.method=='POST':
		try:
			stud_enroll=request.form['stud_enroll']
			stud_name=request.form['stud_name']
			birth_date = request.form["birth_date"] 
			stud_address = request.form["stud_address"] 
			stud_branch = request.form["stud_branch"] 
			stud_branch_code = request.form["stud_branch_code"] 
			stud_marks=request.form["stud_marks"]
			stud_back_sub = request.form["stud_back_sub"] 
			stud_in_comp=request.form["stud_in_comp"] 
			
			with sqlite3.connect("student.db") as con: 
					
				cur = con.cursor()   
				cur.execute("UPDATE student SET stud_name=?,birth_date=?,stud_address=?,stud_branch=?,stud_branch_code=?,stud_marks=?,stud_back_sub=?,stud_in_comp=? WHERE stud_enroll=?",(stud_name,birth_date,stud_address,stud_branch,stud_branch_code,stud_marks,stud_back_sub,stud_in_comp,stud_enroll))  
			
		except:  
			flash("We can not update the data.")
			return redirect('/list5')

		finally:  
			con = sqlite3.connect("student.db")  
			con.row_factory = sqlite3.Row  
			con.commit()
			flash("Data updated successfully!")
			return redirect('/list5')
			

@app.route('/list5')
def list5():
	
	con = sqlite3.connect("student.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from student")   
	details = cur.fetchone()     
	return render_template("add_detail.html",details=details)
	

@app.route("/viewCompany")
def viewCompany():
	title = "List Of Companies"
	con = sqlite3.connect("Company.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from Company")   
	companies = cur.fetchall()  
	totalcomp=len(companies)
	return render_template("viewCompany.html", title=title,companies=companies,totalcomp=totalcomp)


@app.route("/apply_comp",methods = ["POST","GET"])
def apply_comp():
	if request.method == "POST":  
		try:
			comp_name = request.form["comp_name"] 
			stud_enroll=request.form['stud_enroll']
			
			 
			with sqlite3.connect("student.db") as con:  
				cur = con.cursor()   
				cur.execute("UPDATE student SET stud_in_comp=? where stud_enroll=?",(comp_name,stud_enroll))  
				con.commit()  
				flash("Applied to the company!!") 
				return redirect('/list5')
		except:  
			con.rollback()
			flash("You can't apply to the company.Try again!!")
			return redirect('/list5')
		finally:  
			return redirect('/list5')

@app.route("/viewEvents")
def viewEvents():
	title = "Events"
	con = sqlite3.connect("events.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from event")   
	events = cur.fetchall()  
	totalevents=len(events) 
	return render_template("viewEvents.html", title=title,events=events,totalevents=totalevents)

@app.route("/viewDataOfRecruiters")
def viewDataOfRecruiters():
	con = sqlite3.connect("recruit.db")  
	con.row_factory = sqlite3.Row  
	cur = con.cursor()  
	cur.execute("select * from recruit")   
	recruiters = cur.fetchall()
	totalrecruiters=len(recruiters)  
	return render_template("viewDataOfRecruiters.html",recruiters=recruiters,totalrecruiters=totalrecruiters)




@app.route("/changepass_stud")
def changepass_stud():
	return render_template("changepass_stud.html")


@app.route("/changepass_stud_db",methods = ["POST"])
def changepass_stud_db():
	if request.method=='POST':
		try:
			new_pass=request.form['new_pass']
			username = request.form["username"]  
			
			with sqlite3.connect("users.db") as con: 
					
				cur = con.cursor()   
				cur.execute('UPDATE Users SET password=? WHERE username=?', [new_pass, username])
				con.commit()
				flash('Password resetted successfully!!')
			
		except:  
			flash("We can not reset the password.")
			return render_template("changepass_stud.html")

		finally:  
			return render_template("changepass_stud.html")


@app.route("/view_stat")
def view_stat():
	return render_template("view_stat.html")


if __name__ == '__main__':
	app.run(debug=True)

@app.route('/resume-builder', methods=['GET', 'POST'])
def resume_builder():
    if request.method == 'POST':
        # Handle form submission and generate the resume document
        return send_file('resume.docx', as_attachment=True)
    else:
        return render_template('resume_builder.html')


from docx import Document
from docx.shared import Inches

def generate_resume(name, email, phone, education, experience):
    # Create a new Word document
    document = Document()

    # Add the user's name to the document
    document.add_heading(name)

    # Add the user's contact information to the document
    document.add_paragraph('Email: ' + email)
    document.add_paragraph('Phone: ' + phone)

    # Add the user's education history to the document
    document.add_heading('Education')
    for degree in education:
        document.add_paragraph(degree['school'] + ', ' + degree['degree'] + ', ' + degree['year'])

    # Add the user's work experience to the document
    document.add_heading('Experience')
    for job in experience:
        document.add_paragraph(job['company'] + ', ' + job['position'] + ', ' + job['years'])
        document.add_paragraph(job['description'])

    # Save the document to a file
    document.save('resume.docx')
